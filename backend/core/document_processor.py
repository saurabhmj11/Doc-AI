"""
Document Processor (Production Optimized)

Layout-aware chunking for logistics PDFs.

Key improvements:

- Block-level extraction (not flattened text)
- Anchor-aware splitting (Shipper / Consignee / Pickup etc.)
- Sentence-aware overlap
- Page-isolated chunking
- Stable chunk sizes
"""

import uuid
import re
import fitz
from docx import Document
import chardet
from typing import Optional
from dataclasses import dataclass

from config import get_settings
from core.model_loader import get_embedding_model

settings = get_settings()


# ============================================================
# DATA STRUCTURE
# ============================================================

@dataclass
class DocumentChunk:
    chunk_id: str
    document_id: str
    text: str
    page: Optional[int]
    chunk_index: int
    embedding: Optional[list[float]] = None


# ============================================================
# PROCESSOR
# ============================================================

class DocumentProcessor:

    def __init__(self):
        self._embedding_model = None
        self.chunk_size = settings.chunk_size   # use char size
        self.chunk_overlap = settings.chunk_overlap

    @property
    def embedding_model(self):
        if self._embedding_model is None:
            self._embedding_model = get_embedding_model()
        return self._embedding_model


# ============================================================
# PDF PARSER (ELITE FIX)
# ============================================================

    def _parse_pdf(self, file_path):

        pages = []

        with fitz.open(file_path) as doc:

            for page_number, page in enumerate(doc, 1):

                blocks = page.get_text("blocks")

                # SORT by layout position (top-to-bottom, left-to-right)
                blocks.sort(key=lambda b: (b[1], b[0]))

                for block in blocks:

                    text = block[4].strip()

                    if not text:
                        continue

                    # Fix merged headers common in logistics PDFs
                    text = re.sub(
                        r"Shipper\s+Consignee",
                        "Shipper:\nConsignee:",
                        text,
                        flags=re.I
                    )

                    pages.append((page_number, text))

        return pages


# ============================================================
# SMART SPLITTING
# ============================================================

    def _split_sections(self, text):

        anchors = [
            r"\bShipper\b",
            r"\bConsignee\b",
            r"\bPickup\b",
            r"\bDelivery\b",
            r"\bWeight\b",
            r"\bCommodity\b",
            r"\bNotes\b",
            r"\bCarrier\b",
            r"\bBilling\b",
            r"\bFreight\b",
        ]

        pattern = "(" + "|".join(anchors) + ")"

        parts = re.split(pattern, text, flags=re.I)

        if len(parts) > 3:
            merged = []
            for i in range(1, len(parts), 2):
                merged.append(parts[i] + " " + parts[i + 1])
            return merged

        # fallback paragraph split
        return re.split(r"\n{2,}", text)


# ============================================================
# CHUNK CREATION (FIXED)
# ============================================================

    def _create_chunks(self, doc_id, pages):

        chunks = []
        idx = 0

        for page_num, block_text in pages:

            sections = self._split_sections(block_text)

            current = ""

            for section in sections:

                section = section.strip()

                if not section:
                    continue

                # CHAR-based size (stable)
                if len(current) + len(section) > self.chunk_size:

                    if current:

                        chunks.append(
                            DocumentChunk(
                                chunk_id=f"{doc_id}_{idx}",
                                document_id=doc_id,
                                text=current.strip(),
                                page=page_num,
                                chunk_index=idx
                            )
                        )

                        idx += 1

                        # sentence-aware overlap
                        sentences = re.split(r'(?<=[.!?])\s+', current)
                        overlap = " ".join(sentences[-2:])
                        current = overlap + " " + section

                else:
                    current = current + "\n" + section if current else section

            if current.strip():
                chunks.append(
                    DocumentChunk(
                        chunk_id=f"{doc_id}_{idx}",
                        document_id=doc_id,
                        text=current.strip(),
                        page=page_num,
                        chunk_index=idx
                    )
                )
                idx += 1

        return chunks


# ============================================================
# EMBEDDINGS
# ============================================================

    def _generate_embeddings(self, chunks):

        texts = [c.text for c in chunks]

        embeddings = self.embedding_model.encode(
            texts,
            convert_to_numpy=True
        )

        if embeddings is None:
            raise RuntimeError("Embedding generation failed")

        for chunk, emb in zip(chunks, embeddings):
            chunk.embedding = emb.tolist()

        return chunks


# ============================================================
# DOCX / TXT
# ============================================================

    def _parse_docx(self, file_path):

        doc = Document(file_path)
        parts = []

        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)

        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))

        full_text = "\n".join(parts)

        return [(1, full_text)] if full_text else []

    def _parse_txt(self, file_path):

        with open(file_path, "rb") as f:
            raw = f.read()

        enc = chardet.detect(raw).get("encoding", "utf-8") or "utf-8"

        text = raw.decode(enc, errors="replace")

        return [(1, text)] if text.strip() else []


# ============================================================
# MAIN ENTRY
# ============================================================

    def process_file(self, file_path, file_type):

        if file_type == "pdf":
            pages = self._parse_pdf(file_path)
        elif file_type == "docx":
            pages = self._parse_docx(file_path)
        elif file_type == "txt":
            pages = self._parse_txt(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")

        doc_id = str(uuid.uuid4())

        chunks = self._create_chunks(doc_id, pages)

        chunks = self._generate_embeddings(chunks)

        return doc_id, chunks
